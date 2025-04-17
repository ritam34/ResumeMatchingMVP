# from docx import Document

# def convertDocxToText(path):
# 	document = Document(path)
# 	return '\n'.join([para.text for para in document.paragraphs])

from docx import Document
import os
import logging

def convert_docx_to_text(path):
    if not os.path.exists(path) or not path.endswith('.docx'):
        raise FileNotFoundError(f"Invalid .docx file path: {path}")

    try:
        document = Document(path)
        text = []

        for para in document.paragraphs:
            if para.text.strip():  # skip empty lines
                text.append(para.text.strip())

        return '\n'.join(text)

    except Exception as e:
        logging.error(f"Error reading DOCX file '{path}': {e}")
        return ""
